# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

path = r'C:\Users\zhuju\Desktop\赛尔号游戏报告.docx'
doc = Document(path)

# 收集所有段落，找出5.1后面的位置
paras = doc.paragraphs
target_idx = None
for i, p in enumerate(paras):
    if '5.1' in p.text:
        target_idx = i
        break

print(f"找到5.1在段落 {target_idx}")

members_info = [
    "（1）盖亚：战神联盟指挥官，被誉为战斗系之王。盖亚是雷伊的哥哥，拥有强大的战斗能力，是盖亚星球的守护者。盖亚的招牌技能为石破天惊，展现出无与伦比的战斗力量。",
    "（2）雷伊：战神联盟队长，赫尔卡星守护者，雷属性精灵。雷伊是盖亚的弟弟，性格温和但实力强大，是战神联盟的精神领袖。雷伊负责守护赫尔卡星的雷神庙，与盖亚共同维护宇宙秩序。",
    "（3）布莱克：暗影系王者，神秘且强大的精灵。布莱克是暗影系精灵的代表，拥有操控暗影的能力，是战神联盟中的重要成员。",
    "（4）卡修斯：地面系与暗影系双重属性精灵，代号无瑕者。卡修斯是战神联盟中的后起之秀，以其敏捷的身法和强大的战斗力著称。",
    "（5）缪斯：战神联盟唯一的女性成员，负责守护天蛇星。缪斯是超能系精灵，性格温柔但实力不容小觑，是联盟中不可或缺的一员。",
]

rel_title = "5.2 成员关系网"
rel_info = [
    "战神联盟成员之间存在复杂而紧密的关系网络，以下为主要关系梳理：",
    "① 亲属关系：",
    "  - 盖亚与雷伊：兄弟关系（核心关系），盖亚为兄，雷伊为弟，二人共同守护宇宙和平，是战神联盟的核心力量。",
    "  - 盖亚与瑞尔斯：兄弟关系，瑞尔斯为盖亚的另一位亲人，是盖亚的兄弟。",
    "  - 盖亚与米瑞斯：表兄弟关系，米瑞斯是盖亚的表弟，同样是强大的战斗系精灵。",
    "② 组织层级关系：",
    "  - 雷伊担任战神联盟队长/领袖，是团队的精神领袖和核心决策者。",
    "  - 盖亚担任指挥官，是战队的核心战力，负责执行任务和战斗指挥。",
    "  - 布莱克、卡修斯、缪斯为战队成员，各自在联盟中承担不同职责。",
    "③ 宿敌/对手关系：",
    "  - 盖亚与艾辛格：宿敌关系，艾辛格是盖亚的主要对手之一，两人多次交锋。",
    "④ 合作伙伴关系：",
    "  - 斯塔奥、克洛伊等精灵与战神联盟成员保持合作关系，共同维护宇宙秩序。",
    "⑤ 联盟成员关系：",
    "  - 五位成员（盖亚、雷伊、布莱克、卡修斯、缪斯）在战神联盟中并肩作战，彼此信任，共同进退，组成了赛尔号世界观中最具影响力的精灵组织。",
]

# 使用 body.insert_paragraph 在指定位置插入
body = doc.element.body
target_para = doc.paragraphs[target_idx]._element

def insert_para_after(ref_elem, text):
    """在 ref_elem 后面插入一个新段落"""
    new_p = OxmlElement('w:p')
    new_r = OxmlElement('w:r')
    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set(qn('xml:space'), 'preserve')
    new_r.append(new_t)
    new_p.append(new_r)
    ref_elem.addnext(new_p)

# 在5.1段落后插入成员介绍
last_elem = target_para
for info in members_info:
    insert_para_after(last_elem, info)
    last_elem = last_elem.getnext()

# 插入5.2关系网标题
insert_para_after(last_elem, rel_title)
title_elem = last_elem.getnext()
# 加粗标题
r_elem = title_elem.find(qn('w:r'))
if r_elem is not None:
    rPr = OxmlElement('w:rPr')
    b = OxmlElement('w:b')
    rPr.append(b)
    r_elem.insert(0, rPr)
last_elem = title_elem

# 插入关系网内容
for info in rel_info:
    insert_para_after(last_elem, info)
    last_elem = last_elem.getnext()

doc.save(path)
print("报告更新完成！战神联盟成员介绍和关系网已添加到5.1和5.2章节中。")